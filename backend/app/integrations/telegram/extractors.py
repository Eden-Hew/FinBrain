import io
import re
import zipfile
from email import policy
from email.parser import BytesParser
from pathlib import PurePosixPath

from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader

from app.config import get_settings
from app.integrations.ocr.engine import ocr_image, ocr_pdf_pages
from app.integrations.telegram.types import ExtractedContent


class ExtractionError(ValueError):
    pass


ALLOWED_EXTENSIONS = {
    ".txt",
    ".md",
    ".csv",
    ".eml",
    ".pdf",
    ".docx",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".bmp",
    ".tiff",
}
MIME_TYPES = {
    ".txt": {"text/plain", "application/octet-stream"},
    ".md": {"text/markdown", "text/plain", "application/octet-stream"},
    ".csv": {"text/csv", "text/plain", "application/vnd.ms-excel", "application/octet-stream"},
    ".eml": {"message/rfc822", "text/plain", "application/octet-stream"},
    ".pdf": {"application/pdf", "application/octet-stream"},
    ".docx": {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/octet-stream",
    },
    ".png": {"image/png", "application/octet-stream"},
    ".jpg": {"image/jpeg", "application/octet-stream"},
    ".jpeg": {"image/jpeg", "application/octet-stream"},
    ".webp": {"image/webp", "application/octet-stream"},
    ".bmp": {"image/bmp", "application/octet-stream"},
    ".tiff": {"image/tiff", "application/octet-stream"},
}
TEXT_EXTENSIONS = {".txt", ".md", ".csv"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"}


def _normalize(text: str) -> str:
    settings = get_settings()
    text = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{4,}", "\n\n\n", text).strip()
    if not text:
        raise ExtractionError("no_extractable_text")
    if len(text) > settings.telegram_max_extracted_chars:
        raise ExtractionError("extracted_text_too_large")
    return text


def _extension(filename: str) -> str:
    suffix = PurePosixPath(filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise ExtractionError("unsupported_file_type")
    return suffix


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "windows-1252"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ExtractionError("malformed_document")


def _extract_eml(data: bytes) -> str:
    try:
        message = BytesParser(policy=policy.default).parsebytes(data)
        fields = [
            f"{label}: {message.get(header, '')}"
            for label, header in (
                ("Subject", "subject"),
                ("From", "from"),
                ("To", "to"),
                ("Date", "date"),
            )
            if message.get(header)
        ]
        plain: list[str] = []
        html: list[str] = []
        parts = list(message.walk()) if message.is_multipart() else [message]
        if len(parts) > 200:
            raise ExtractionError("document_expansion_limit")
        for part in parts:
            if part.get_content_disposition() == "attachment":
                continue
            content_type = part.get_content_type()
            if content_type not in {"text/plain", "text/html"}:
                continue
            content = part.get_content()
            if content_type == "text/plain":
                plain.append(str(content))
            else:
                soup = BeautifulSoup(str(content), "html.parser")
                for element in soup(["script", "style", "form", "noscript"]):
                    element.decompose()
                html.append(soup.get_text("\n"))
        body = "\n".join(plain or html)
        return "\n".join(fields + ([""] if fields and body else []) + [body])
    except ExtractionError:
        raise
    except Exception as error:
        raise ExtractionError("malformed_document") from error


def _extract_pdf(data: bytes) -> tuple[str, int, str]:
    if not data.startswith(b"%PDF-"):
        raise ExtractionError("invalid_file_signature")
    try:
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            raise ExtractionError("encrypted_pdf")
        if len(reader.pages) > get_settings().telegram_max_pdf_pages:
            raise ExtractionError("pdf_page_limit")
        page_texts = [page.extract_text() or "" for page in reader.pages]
        raw_text = "\n".join(text.strip() for text in page_texts if text.strip())
        settings = get_settings()
        if settings.enable_ocr and len(raw_text) < settings.ocr_min_text_chars:
            ocr_pages = ocr_pdf_pages(data, max_pages=settings.ocr_max_pages)
            if any(text.strip() for text in ocr_pages):
                return (
                    "\n\n".join(
                        f"--- Page {index} ---\n{text}"
                        for index, text in enumerate(ocr_pages, 1)
                        if text.strip()
                    ),
                    len(ocr_pages),
                    "ocr",
                )
        if not raw_text:
            raise ExtractionError("no_extractable_text")
        pages = [
            f"--- Page {index} ---\n{text}" for index, text in enumerate(page_texts, 1)
        ]
        return "\n\n".join(pages), len(reader.pages), "text"
    except ExtractionError:
        raise
    except Exception as error:
        raise ExtractionError("malformed_document") from error


def _extract_image(data: bytes, suffix: str) -> str:
    if len(data) > get_settings().ocr_max_image_bytes:
        raise ExtractionError("file_too_large")
    text = ocr_image(data)
    if not text:
        raise ExtractionError("no_extractable_text")
    return text


def _validate_docx_zip(data: bytes) -> None:
    if not data.startswith(b"PK"):
        raise ExtractionError("invalid_file_signature")
    settings = get_settings()
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            members = archive.infolist()
            if len(members) > settings.telegram_max_docx_members:
                raise ExtractionError("document_expansion_limit")
            expanded_size = sum(member.file_size for member in members)
            if expanded_size > settings.telegram_max_docx_uncompressed_bytes:
                raise ExtractionError("document_expansion_limit")
            for member in members:
                path = PurePosixPath(member.filename)
                if path.is_absolute() or ".." in path.parts:
                    raise ExtractionError("malformed_document")
            if "word/document.xml" not in archive.namelist():
                raise ExtractionError("malformed_document")
    except ExtractionError:
        raise
    except Exception as error:
        raise ExtractionError("malformed_document") from error


def _extract_docx(data: bytes) -> str:
    _validate_docx_zip(data)
    try:
        document = Document(io.BytesIO(data))
        chunks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                chunks.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(chunks)
    except Exception as error:
        raise ExtractionError("malformed_document") from error


def extract_document(data: bytes, *, filename: str, mime_type: str) -> ExtractedContent:
    settings = get_settings()
    if len(data) > settings.telegram_max_file_bytes:
        raise ExtractionError("file_too_large")
    suffix = _extension(filename)
    normalized_mime = (mime_type or "application/octet-stream").lower()
    if normalized_mime not in MIME_TYPES[suffix]:
        raise ExtractionError("unsupported_file_type")
    page_count = None
    extraction_method = None
    if suffix in TEXT_EXTENSIONS:
        text = _decode_text(data)
        extraction_method = "text"
    elif suffix == ".eml":
        text = _extract_eml(data)
        extraction_method = "text"
    elif suffix == ".pdf":
        text, page_count, extraction_method = _extract_pdf(data)
    elif suffix in IMAGE_EXTENSIONS:
        text = _extract_image(data, suffix)
        extraction_method = "ocr"
    else:
        text = _extract_docx(data)
        extraction_method = "text"
    return ExtractedContent(
        text=_normalize(text),
        input_kind=suffix.removeprefix("."),
        mime_type=normalized_mime,
        filename=filename,
        page_count=page_count,
        extraction_method=extraction_method,
    )


def extract_text(text: str, *, forwarded: bool = False) -> ExtractedContent:
    return ExtractedContent(
        text=_normalize(text),
        input_kind="forwarded_text" if forwarded else "text",
        mime_type="text/plain",
    )
